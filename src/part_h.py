#!/usr/bin/env python3
"""Part H: Comparison of Image Descriptions (10 points)

In this part, we will test the quality of the different representations. A good representation 
is one that retains some of the semantics of the image; oftentimes by "semantics" we mean object 
class label. In other words, a good representation should be one such that two images of the same 
object have similar representations, and images of different objects have different representations.

To test the quality of the representations, we will compare two averages: the average within-class 
distance and the average between-class distance. A representation is good if the "distance" is the 
Euclidean distance between two vectors (i.e. the representations of two images). "Within-class 
distances" are distances computed between the vectors for images of the same class (i.e. cardinal-cardinal, 
panda-panda). "Between-class distances" are those computed between images of different classes, i.e. 
cardinal-panda, panda-leopard, etc.

Usage example:
    python src/part_h.py --data-dir data --out-dir out
"""
import os
import argparse
import numpy as np
import pickle
from docx import Document
from datetime import datetime

from utils import load_and_preprocess_image, ensure_dir

EXPECTED_IMAGES = [
    ("cardinal1", "cardinal1.jpg"),
    ("cardinal2", "cardinal2.jpg"), 
    ("leopard1", "leopard1.jpg"),
    ("leopard2", "leopard2.jpg"),
    ("panda1", "panda1.jpg"),
    ("panda2", "panda2.jpg"),
]

# Define animal categories for within/between class comparison
CATEGORIES = {
    'cardinal': ['cardinal1', 'cardinal2'],
    'leopard': ['leopard1', 'leopard2'],
    'panda': ['panda1', 'panda2']
}


def load_bow_representations(part_e_dir):
    """Load SIFT bag-of-words representations from Part E."""
    npz_path = os.path.join(part_e_dir, 'bow_histograms.npz')
    if not os.path.exists(npz_path):
        raise FileNotFoundError(f"BoW representations not found at {npz_path}. Run Part E first.")
    
    data = np.load(npz_path)
    histograms = data['histograms']
    image_names = data['image_names']
    
    bow_repr = {}
    for i, name in enumerate(image_names):
        bow_repr[name] = histograms[i]
    
    return bow_repr


def load_texture_representations(part_b_dir):
    """Load texture representations from Part B."""
    npz_path = os.path.join(part_b_dir, 'texture_representations.npz') 
    if not os.path.exists(npz_path):
        raise FileNotFoundError(f"Texture representations not found at {npz_path}. Run Part B first.")
    
    data = np.load(npz_path)
    
    texture_repr_concat = {}
    for category in ['cardinal', 'leopard', 'panda']:
        key = f'{category}_concat'
        if key in data:
            for img_name in CATEGORIES[category]:
                texture_repr_concat[img_name] = data[key]
    
    texture_repr_mean = {}
    for category in ['cardinal', 'leopard', 'panda']:
        key = f'{category}_mean'
        if key in data:
            for img_name in CATEGORIES[category]:
                texture_repr_mean[img_name] = data[key]
    
    return texture_repr_concat, texture_repr_mean


def compute_within_between_distances(representations, image_names):
    """Compute within-class and between-class distances for given representations.
    
    Args:
        representations: dict mapping image_name -> representation vector
        image_names: list of image names to analyze
        
    Returns:
        within_distances: list of within-class distances
        between_distances: list of between-class distances
        avg_within: average within-class distance
        avg_between: average between-class distance
        ratio: within/between ratio (as specified in Part H)
    """
    within_distances = []
    between_distances = []
    
    for i, name1 in enumerate(image_names):
        for j, name2 in enumerate(image_names):
            if i >= j:
                continue
                
            if name1 not in representations or name2 not in representations:
                continue
                
            vec1 = representations[name1]
            vec2 = representations[name2]
            distance = np.linalg.norm(vec1 - vec2)
            
            category1 = None
            category2 = None
            
            for cat, images in CATEGORIES.items():
                if name1 in images:
                    category1 = cat
                if name2 in images:
                    category2 = cat
            
            if category1 == category2 and category1 is not None:
                within_distances.append(distance)
            elif category1 != category2 and category1 is not None and category2 is not None:
                between_distances.append(distance)
    
    avg_within = np.mean(within_distances) if within_distances else 0.0
    avg_between = np.mean(between_distances) if between_distances else 0.0
    
    ratio = avg_within / avg_between if avg_between > 0 else float('inf')
    
    return within_distances, between_distances, avg_within, avg_between, ratio


def create_word_document(results, save_path):
    """Create a Word document with the Part H results."""
    doc = Document()
    
    title = doc.add_heading('Part H: Comparison of Image Descriptions', 0)
    title.alignment = 1
    
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = 1
    
    doc.add_paragraph()
    
    intro = doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This document presents the comparison of three different image representation methods '
        'by analyzing within-class and between-class distances. A good representation should '
        'have smaller within-class distances (images of same object are similar) and larger '
        'between-class distances (images of different objects are different).'
    )
    
    method = doc.add_heading('Methodology', level=1)
    doc.add_paragraph('• Dataset: 6 animal images (2 cardinals, 2 leopards, 2 pandas)')
    doc.add_paragraph('• Distance metric: Euclidean distance between representation vectors')
    doc.add_paragraph('• Within-class distances: Between images of same animal category')
    doc.add_paragraph('• Between-class distances: Between images of different animal categories')
    doc.add_paragraph('• Quality metric: average_within_class_distance / average_between_class_distance')
    doc.add_paragraph('• Lower ratio indicates better discriminative power')
    
    summary = doc.add_heading('Results Summary', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Representation Method'
    hdr_cells[1].text = 'Avg Within-Class'
    hdr_cells[2].text = 'Avg Between-Class'
    hdr_cells[3].text = 'Within/Between Ratio'
    hdr_cells[4].text = 'Rank'
    
    sorted_results = sorted(results.items(), key=lambda x: x[1]['ratio'])
    
    for rank, (method, data) in enumerate(sorted_results, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = method.replace('_', ' ').title()
        row_cells[1].text = f"{data['avg_within']:.6f}"
        row_cells[2].text = f"{data['avg_between']:.6f}"
        row_cells[3].text = f"{data['ratio']:.6f}"
        row_cells[4].text = str(rank)
    
    details = doc.add_heading('Detailed Analysis', level=1)
    
    for method, data in sorted_results:
        method_heading = doc.add_heading(method.replace('_', ' ').title(), level=2)
        
        doc.add_paragraph(f"Average within-class distance: {data['avg_within']:.6f}")
        doc.add_paragraph(f"Average between-class distance: {data['avg_between']:.6f}")
        doc.add_paragraph(f"Within/between ratio: {data['ratio']:.6f}")
        doc.add_paragraph(f"Number of within-class pairs: {len(data['within_distances'])}")
        doc.add_paragraph(f"Number of between-class pairs: {len(data['between_distances'])}")
        
        if data['within_distances']:
            within_range = f"[{min(data['within_distances']):.4f}, {max(data['within_distances']):.4f}]"
            doc.add_paragraph(f"Within-class distance range: {within_range}")
        
        if data['between_distances']:
            between_range = f"[{min(data['between_distances']):.4f}, {max(data['between_distances']):.4f}]"
            doc.add_paragraph(f"Between-class distance range: {between_range}")
        
        doc.add_paragraph()
    
    conclusions = doc.add_heading('Conclusions', level=1)
    best_method = sorted_results[0][0].replace('_', ' ').title()
    best_ratio = sorted_results[0][1]['ratio']
    
    doc.add_paragraph(
        f'The {best_method} method achieved the best performance with a within/between ratio '
        f'of {best_ratio:.6f}. This indicates that this representation method provides the '
        f'best discriminative power for distinguishing between different animal categories '
        f'while maintaining similarity within the same category.'
    )
    
    doc.save(save_path)


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--data-dir', default='data')
    p.add_argument('--out-dir', default='out')
    args = p.parse_args()

    out_dir = args.out_dir
    ensure_dir(out_dir)
    ensure_dir(os.path.join(out_dir, 'part_h'))
    
    save_dir = os.path.join(out_dir, 'part_h')
    part_b_dir = os.path.join(out_dir, 'part_b')
    part_e_dir = os.path.join(out_dir, 'part_e')

    print("Part H: Comparison of Image Descriptions (10 points)")
    print("=" * 55)
    
    print("Loading representations...")
    
    try:
        image_names = [name for name, _ in EXPECTED_IMAGES]
        print(f"✓ Image list prepared: {len(image_names)} images")
        
        bow_repr = load_bow_representations(part_e_dir)
        print(f"✓ Loaded BoW representations: {len(bow_repr)} images")
        
        texture_repr_concat, texture_repr_mean = load_texture_representations(part_b_dir)
        print(f"✓ Loaded texture concat representations: {len(texture_repr_concat)} images")
        print(f"✓ Loaded texture mean representations: {len(texture_repr_mean)} images")
        
    except FileNotFoundError as e:
        print(f"Error: {e}")
        print("Please run Parts B and E first to generate the required representations.")
        return

    representations = {
        'bow_repr': bow_repr,
        'texture_repr_concat': texture_repr_concat, 
        'texture_repr_mean': texture_repr_mean
    }
    
    results = {}
    
    print(f"\nComputing within-class and between-class distances...")
    
    for repr_name, repr_data in representations.items():
        print(f"\nAnalyzing {repr_name}:")
        
        within_dist, between_dist, avg_within, avg_between, ratio = compute_within_between_distances(
            repr_data, image_names
        )
        
        results[repr_name] = {
            'within_distances': within_dist,
            'between_distances': between_dist,
            'avg_within': avg_within,
            'avg_between': avg_between,
            'ratio': ratio
        }
        
        print(f"  Average within-class distance: {avg_within:.6f}")
        print(f"  Average between-class distance: {avg_between:.6f}")
        print(f"  Ratio (average_within_class_distance / average_between_class_distance): {ratio:.6f}")
        print(f"  Within-class pairs: {len(within_dist)}")
        print(f"  Between-class pairs: {len(between_dist)}")
    
    print(f"\n" + "=" * 55)
    print("SUMMARY COMPARISON")
    print("=" * 55)
    
    sorted_results = sorted(results.items(), key=lambda x: x[1]['ratio'])
    
    print(f"{'Method':<25} {'Within':<10} {'Between':<10} {'Ratio':<10} {'Rank':<5}")
    print("-" * 65)
    
    for rank, (method, data) in enumerate(sorted_results, 1):
        print(f"{method.replace('_', ' '):<25} {data['avg_within']:<10.4f} "
              f"{data['avg_between']:<10.4f} {data['ratio']:<10.4f} {rank:<5}")
    
    best_method = sorted_results[0][0]
    print(f"\nBest performing method: {best_method.replace('_', ' ')} "
          f"(ratio: {sorted_results[0][1]['ratio']:.6f})")
    
    results_file = os.path.join(save_dir, 'part_h_results.npz')
    np.savez(results_file, **{f'{k}_data': v for k, v in results.items()})
    print(f"\nSaved numerical results: {results_file}")
    
    print(f"\nCreating Word document...")
    word_file = os.path.join(save_dir, 'results.doc')
    
    try:
        create_word_document(results, word_file)
        print(f"✓ Created Word document: {word_file}")
        print(f"  Document contains ratios for all three representation methods")
        print(f"  Lower ratio indicates better discriminative power")
    except ImportError:
        print(f"Warning: python-docx not installed. Creating text file instead.")
        
        text_file = os.path.join(save_dir, 'results.txt')
        with open(text_file, 'w') as f:
            f.write("Part H: Comparison of Image Descriptions\n")
            f.write("=" * 45 + "\n\n")
            
            for method, data in sorted_results:
                f.write(f"{method.replace('_', ' ').title()}:\n")
                f.write(f"  Average within-class distance: {data['avg_within']:.6f}\n")
                f.write(f"  Average between-class distance: {data['avg_between']:.6f}\n")
                f.write(f"  Ratio (within/between): {data['ratio']:.6f}\n\n")
        
        print(f"✓ Created text file: {text_file}")

    print(f"\n" + "=" * 55)
    print("Part H completed successfully!")
    print("Key findings:")
    print(f"• {best_method.replace('_', ' ')} provides best discrimination")
    print(f"• Lower within/between ratio indicates better representation quality")
    print(f"• Results saved in both numerical and document formats")


if __name__ == '__main__':
    main()